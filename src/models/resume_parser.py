import os
from docx import Document
import PyPDF2
import re
from src.config import regex_config

class ResumeParser:

    def read_docx(self, file_path):
        doc = Document(file_path)
        text = ""
        
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text.strip() + "\n\n"
        
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
            text += "\n"
        
        for section in ["EDUCATION", "EXPERIENCE", "SKILLS", "WORK EXPERIENCE", "PROJECTS", "TECHNICAL SKILLS"]:
            pattern = re.compile(f"({section})", re.IGNORECASE)
            text = pattern.sub(f"\n\n{section.upper()}\n", text)
            
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        print(f"\nExtraído del documento DOCX (primeros 500 caracteres):\n{text[:500]}...\n")
        
        return text

    def read_pdf(self, file_path):
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            
            spaced_text_pattern = r'\b([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\b'
            text = re.sub(spaced_text_pattern, r'\1\2\3\4\5\6\7\8', text)
            spaced_text_pattern = r'\b([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\b'
            text = re.sub(spaced_text_pattern, r'\1\2\3\4\5\6\7', text)
            spaced_text_pattern = r'\b([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\b'
            text = re.sub(spaced_text_pattern, r'\1\2\3\4\5\6', text)
            spaced_text_pattern = r'\b([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\b'
            text = re.sub(spaced_text_pattern, r'\1\2\3\4\5', text)
            spaced_text_pattern = r'\b([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\b'
            text = re.sub(spaced_text_pattern, r'\1\2\3\4', text)
            
            spaced_text_pattern = r'\b([A-Z])\s+([A-Z])\s+([A-Z])\b'
            text = re.sub(spaced_text_pattern, r'\1\2\3', text)
            spaced_text_pattern = r'\b([A-Z])\s+([A-Z])\b'
            text = re.sub(spaced_text_pattern, r'\1\2', text)
            
            spaced_word_pattern = r'\b([A-Za-z])\s+([a-z])\s+([a-z])\b'
            text = re.sub(spaced_word_pattern, r'\1\2\3', text)
            
            return text

    def read_txt(self, file_path):
        with open(file_path, 'r') as file:
            return file.read()

    def read_resume(self, file_path):
        extension = os.path.splitext(file_path)[1].lower()  
        if extension == '.docx':
            return self.read_docx(file_path)
        elif extension == '.pdf':
            return self.read_pdf(file_path)
        elif extension == '.txt':
            return self.read_txt(file_path)
        else:
            raise ValueError("Unsupported file type. Please upload a .docx, .pdf, or .txt file.")

    def extract_resume_info(self, resume_text, role):
        info = {}

        print(f"Analizando texto (primeros 100 caracteres): {resume_text[:100]}")
        
        role_default_skills = {
            "Data Scientist": "Python, R, SQL, Machine Learning, Data Analysis, Statistics, Data Visualization",
            "AI Engineer": "Python, TensorFlow, PyTorch, Computer Vision, NLP, Deep Learning",
            "Software Engineer": "Java, Python, C++, OOP, Algorithms, Data Structures, Git",
            "Web Developer": "HTML, CSS, JavaScript, React, Node.js, RESTful APIs, Responsive Design",
            "DevOps Engineer": "Docker, Kubernetes, CI/CD, AWS, Linux, Scripting, Infrastructure as Code"
        }
        
        name_patterns = [
            r"^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)",
            r"^([A-Z][A-Z\s]+)(?:\n|$)",
            r"^([A-Z][a-z]+\s+[A-Za-z\-']+)",
            r"([A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)"
        ]
        
        first_lines = resume_text.split('\n')[:3]
        first_text = '\n'.join(first_lines)
        
        name_found = False
        
        for pattern in name_patterns:
            name_match = re.search(pattern, first_text)
            if name_match and name_match.group(1).strip():
                candidate = name_match.group(1).strip()
                
                if (len(candidate.split()) >= 2
                    and not any(section in candidate.upper() for section in ["EDUCATION", "EXPERIENCE", "SKILLS", "PROJECTS", "WORK"])
                    and not re.search(r',\s*[A-Z]{2}', candidate)
                    and len(candidate) <= 40
                    and not re.match(r'.*\d+.*', candidate)
                    and ":" not in candidate
                    and "|" not in candidate
                    and not re.search(r'(University|College|School|Institute|degree|resume|cv)', candidate.lower())):
                    
                    info['name'] = candidate
                    name_found = True
                    break
        
        if not name_found:
            clean_first_line = first_lines[0].strip() if first_lines else ""
            if (clean_first_line and len(clean_first_line) <= 40 
                and not re.search(r'[,@:|/\\"]', clean_first_line)
                and not re.search(r',\s*[A-Z]{2}', clean_first_line)
                and not re.match(r'.*\d+.*', clean_first_line)
                and all(word[0].isupper() for word in clean_first_line.split() if word)):
                
                info['name'] = clean_first_line
        
        if 'name' not in info or not info['name']:
            name_header_patterns = [
                r'Name:\s*([A-Za-z\s\-\']+)',
                r'Resume of\s*([A-Za-z\s\-\']+)',
                r'CV of\s*([A-Za-z\s\-\']+)',
                r'Curriculum Vitae of\s*([A-Za-z\s\-\']+)'
            ]
            
            for pattern in name_header_patterns:
                match = re.search(pattern, resume_text[:500], re.IGNORECASE)
                if match and match.group(1).strip():
                    info['name'] = match.group(1).strip()
                    break
        
        if 'name' not in info or not info['name']:
            info['name'] = "John Doe"
            
        email_patterns = [
            regex_config.email_regex,
            r'[\w\.-]+@[\w\.-]+\.[\w]{2,}',
            r'[\w\.-]+@[\w\.-]+',
            r'(?:email|e-mail|Email|E-mail)[\s:]*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
            r'(?:E:|P:|T:)?\s*[\w\.-]+@[\w\.-]+\.[\w]{2,}'
        ]
    
        for pattern in email_patterns:
            email_match = re.search(pattern, resume_text, re.IGNORECASE)
            if email_match:
                email_candidate = email_match.group(0).strip()
                
                if '@' in email_candidate and '.' in email_candidate.split('@')[1]:
                    info['email'] = email_candidate
                    break
                
        if 'email' not in info or not info['email']:
            info['email'] = "example@email.com"
            
        phone_patterns = [
            r'P:\s*(\+\d{1,3}[\s.-]?\d{1,9})',
            r'(\+\d{1,3}[\s.-]?\d{1,3}[\s.-]?\d{4,10})',
            regex_config.phone_regex,
            r'(\+\d{1,3}[\s.-]?\d{1,9})',
            r'(\+?\d{1,3}[-.\s]??\d{3}[-.\s]??\d{3}[-.\s]??\d{4})',
            r'(\d{3}[-.\s]??\d{3}[-.\s]??\d{4})',
            r'(\(\d{3}\)\s*\d{3}[-.\s]??\d{4})',
            r'(\d{10})',
            r'(?:phone|mobile|cell|Phone|Mobile|Cell|P:)[\s:]*([0-9()+\-.\s]{7,})'
        ]
    
        for pattern in phone_patterns:
            phone_match = re.search(pattern, resume_text)
            if phone_match:
                phone = phone_match.group(0).strip()
                if phone.startswith("P:"):
                    phone = phone[2:].strip()
                    
                if re.search(r'\d', phone):
                    if '@' not in phone and not re.match(r'[\w\.-]+@[\w\.-]+', phone):
                        info['phone'] = phone
                        break
                
        if 'phone' not in info or not info['phone']:
            info['phone'] = "+1-234-567-8901"
            
        location_found = False
        
        first_lines = resume_text.split('\n')[:5]
        
        for line in first_lines:
            location_match = re.search(r'([A-Za-z\s]+,\s*(?:AL|AK|AZ|AR|CA|CO|CT|DE|FL|GA|HI|ID|IL|IN|IA|KS|KY|LA|ME|MD|MA|MI|MN|MS|MO|MT|NE|NV|NH|NJ|NM|NY|NC|ND|OH|OK|OR|PA|RI|SC|SD|TN|TX|UT|VT|VA|WA|WV|WI|WY))', line)
            if location_match:
                candidate = location_match.group(1).strip()
                if 'name' in info and candidate.upper() != info['name'].upper():
                    info['location'] = candidate
                    location_found = True
                    break
        
        if not location_found:
            contact_line_pattern = r'(?:^|\n)([^|]+\|\s*[^|]+\|\s*[^|]+)(?:$|\n)'
            contact_line_match = re.search(contact_line_pattern, resume_text[:500])
            if contact_line_match:
                contact_line = contact_line_match.group(1)
                parts = [p.strip() for p in contact_line.split('|')]
                for part in parts:
                    if re.search(r'[A-Za-z\s]+,\s*[A-Z]{2}', part) and 'email' not in part.lower() and '@' not in part:
                        info['location'] = part
                        location_found = True
                        break
        
        if not location_found:
            location_patterns = [
                r'Location[:\s]+([A-Za-z\s]+,\s*[A-Za-z\s]+)',
                r'Address[:\s]+([A-Za-z\s]+,\s*[A-Za-z\s]+)',
                r'Based in[:\s]+([A-Za-z\s]+,\s*[A-Za-z\s]+)',
                r'(?<!\w)([A-Za-z\s]+,\s*(?:AL|AK|AZ|AR|CA|CO|CT|DE|FL|GA|HI|ID|IL|IN|IA|KS|KY|LA|ME|MD|MA|MI|MN|MS|MO|MT|NE|NV|NH|NJ|NM|NY|NC|ND|OH|OK|OR|PA|RI|SC|SD|TN|TX|UT|VT|VA|WA|WV|WI|WY))(?!\w)',
                r'(?<!\w)([A-Za-z\s]+,\s*[A-Z]{2})(?!\w)',
            ]
            
            location_section_matches = re.finditer(r'(?:Location|Address|Contact)(?::|\s+)([^\n|]+)', resume_text[:500], re.IGNORECASE)
            for match in location_section_matches:
                location = match.group(1).strip()
                if len(location) > 3 and re.search(r'[A-Za-z\s]+,\s*[A-Za-z]{2}', location):
                    info['location'] = location
                    location_found = True
                    break
            
            if not location_found:
                for pattern in location_patterns:
                    location_match = re.search(pattern, resume_text[:500], re.IGNORECASE)
                    if location_match and location_match.group(1).strip():
                        candidate = location_match.group(1).strip()
                        
                        if ('name' in info and candidate.upper() == info['name'].upper()):
                            continue
                        
                        if (not any(section in candidate.upper() for section in ["EDUCATION", "EXPERIENCE", "SKILLS", "PROJECTS", "WORK"]) 
                            and not candidate == info.get('name', '')
                            and not '@' in candidate
                            and not re.search(r'^\+?\d', candidate)
                            and not "assign tasks" in candidate
                            and not "WORK EXPERIENCE" in candidate):
                            info['location'] = candidate
                            location_found = True
                            break
        
        if not location_found:
            common_locations = [
                "New York, NY", "San Francisco, CA", "Los Angeles, CA", "Chicago, IL", 
                "Seattle, WA", "Boston, MA", "Austin, TX", "Atlanta, GA", "Denver, CO",
                "Tampa, FL", "Miami, FL", "Dallas, TX", "Houston, TX", "Portland, OR"
            ]
            
            for location in common_locations:
                if location.lower() in resume_text.lower():
                    info['location'] = location
                    location_found = True
                    break
                    
        if not location_found or ('name' in info and info['location'].upper() == info['name'].upper()):
            info['location'] = "Remote"
        
        if "linkedin" in resume_text.lower() or "linked in" in resume_text.lower():
            first_hundred_lines = '\n'.join(resume_text.split('\n')[:15])
            
            linkedin_url_patterns = [
                r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\.-]+',
                r'(?:https?://)?(?:www\.)?linkedin\.com/[\w\.-]+'
            ]
            
            for pattern in linkedin_url_patterns:
                url_match = re.search(pattern, first_hundred_lines, re.IGNORECASE)
                if url_match:
                    linkedin_url = url_match.group(0).strip()
                    if not linkedin_url.startswith('http'):
                        linkedin_url = f"https://{linkedin_url}"
                    info['linkedin'] = linkedin_url
                    break
                    
            if 'linkedin' not in info or not info['linkedin']:
                lines = first_hundred_lines.split('\n')
                for i, line in enumerate(lines):
                    if ('linkedin' in line.lower() or 'linked in' in line.lower()) and len(line) < 100:
                        words = re.split(r'[\s\|\(\)\[\]\{\}]', line)
                        
                        for word in words:
                            word = word.strip('.,;:')
                            if len(word) > 3 and word.lower() != 'linkedin' and not re.match(r'linked\s*in', word.lower()):
                                if all(c.isalnum() or c in '-_.' for c in word):
                                    if 'linkedin.com' in word:
                                        if not word.startswith('http'):
                                            info['linkedin'] = f"https://{word}"
                                        else:
                                            info['linkedin'] = word
                                    else:
                                        info['linkedin'] = f"https://linkedin.com/in/{word}"
                                    break
                        
                        if 'linkedin' in info and info['linkedin']:
                            break
                            
                        if i < len(lines) - 1:
                            next_line = lines[i + 1].strip()
                            if len(next_line) < 100:
                                words = re.split(r'[\s\|\(\)\[\]\{\}]', next_line)
                                for word in words:
                                    word = word.strip('.,;:')
                                    if len(word) > 3 and 'linkedin.com' in word:
                                        if not word.startswith('http'):
                                            info['linkedin'] = f"https://{word}"
                                        else:
                                            info['linkedin'] = word
                                        break
                                    elif len(word) > 3 and all(c.isalnum() or c in '-_.' for c in word):
                                        info['linkedin'] = f"https://linkedin.com/in/{word}"
                                        break
                
            if 'linkedin' not in info or not info['linkedin']:
                if 'name' in info and info['name']:
                    name_parts = info['name'].split()
                    if len(name_parts) >= 2:
                        first_name = name_parts[0].lower()
                        last_name = name_parts[-1].lower()
                        
                        potential_usernames = [
                            f"{first_name}-{last_name}",
                            f"{first_name}{last_name}",
                            f"{first_name}.{last_name}",
                            f"{first_name[0]}{last_name}"
                        ]
                        
                        for username in potential_usernames:
                            if re.search(username, resume_text, re.IGNORECASE):
                                info['linkedin'] = f"https://linkedin.com/in/{username}"
                                break
        
        if 'linkedin' not in info or not info['linkedin'] or 'linkedin profile' in info['linkedin'].lower():
            info['linkedin'] = "https://linkedin.com/in/johndoe"
        
        if "github" in resume_text.lower() or "portfolio" in resume_text.lower() or "web" in resume_text.lower():
            portfolio_url_patterns = [
                r'(?:https?://)?(?:www\.)?github\.com/[\w\.-]+',
                r'(?:https?://)?(?:www\.)?[\w\.-]+\.(?:com|io|dev|net|org)(?:/[\w\.-]*)?',
                r'(?:https?://)?[\w\.-]+\.(?:com|io|dev|net|org)(?:/[\w\.-]*)?'
            ]
            
            for pattern in portfolio_url_patterns:
                url_match = re.search(pattern, resume_text, re.IGNORECASE)
                if url_match:
                    portfolio_url = url_match.group(0).strip()
                    if not portfolio_url.startswith('http'):
                        portfolio_url = f"https://{portfolio_url}"
                    info['portfolio'] = portfolio_url
                    break
                    
            if 'portfolio' not in info or not info['portfolio']:
                first_hundred_lines = '\n'.join(resume_text.split('\n')[:15])
                lines = first_hundred_lines.split('\n')
                
                for i, line in enumerate(lines):
                    if ('github' in line.lower() or 'portfolio' in line.lower() or 'website' in line.lower() or 'web' in line.lower()) and len(line) < 100:
                        words = re.split(r'[\s\|\(\)\[\]\{\}]', line)
                        
                        for word in words:
                            word = word.strip('.,;:')
                            if len(word) > 3 and word.lower() != 'github' and word.lower() != 'portfolio' and word.lower() != 'website':
                                if '.' in word:
                                    if not word.startswith('http'):
                                        info['portfolio'] = f"https://{word}"
                                    else:
                                        info['portfolio'] = word
                                    break
                                elif all(c.isalnum() or c in '-_' for c in word):
                                    info['portfolio'] = f"https://github.com/{word}"
                                    break
                        
                        if 'portfolio' in info and info['portfolio']:
                            break
                            
                        if i < len(lines) - 1:
                            next_line = lines[i + 1].strip()
                            if len(next_line) < 100:
                                words = re.split(r'[\s\|\(\)\[\]\{\}]', next_line)
                                for word in words:
                                    word = word.strip('.,;:')
                                    if len(word) > 3:
                                        if '.' in word:
                                            if not word.startswith('http'):
                                                info['portfolio'] = f"https://{word}"
                                            else:
                                                info['portfolio'] = word
                                            break
                                        elif all(c.isalnum() or c in '-_' for c in word):
                                            info['portfolio'] = f"https://github.com/{word}"
                                            break
                
            if 'portfolio' not in info or not info['portfolio']:
                if 'name' in info and info['name']:
                    name_parts = info['name'].split()
                    if len(name_parts) >= 1:
                        first_name = name_parts[0].lower()
                        last_name = name_parts[-1].lower() if len(name_parts) > 1 else ""
                        
                        portfolio_patterns = [
                            rf'(?:https?://)?(?:www\.)?{first_name}[\w\.-]*\.(?:com|io|dev|net|org)',
                            rf'(?:https?://)?(?:www\.)?{last_name}[\w\.-]*\.(?:com|io|dev|net|org)' if last_name else None,
                            rf'(?:https?://)?(?:www\.)?{first_name}-{last_name}\.(?:com|io|dev|net|org)' if last_name else None,
                            rf'(?:https?://)?(?:www\.)?{first_name}{last_name}\.(?:com|io|dev|net|org)' if last_name else None
                        ]
                        
                        for pattern in portfolio_patterns:
                            if pattern:
                                match = re.search(pattern, resume_text, re.IGNORECASE)
                                if match:
                                    url = match.group(0)
                                    if not url.startswith('http'):
                                        url = f"https://{url}"
                                    info['portfolio'] = url
                                    break
        
        if 'portfolio' not in info or not info['portfolio'] or any(invalid in info['portfolio'].lower() for invalid in ['portfolio', 'github\nwork']):
            info['portfolio'] = "https://github.com/example"
            
        education_patterns = [
            regex_config.education_regex,
            r'EDUCATION.*?\n(.*?)(?=\n\s*\n|\n[A-Z\s]{2,}|\Z)',
            r'(?:University|College|Institute).*?(?:Bachelor|Master|PhD|BS|MS|BA|Degree)',
            r'(?:BS|MS|BA|PhD|Bachelor|Master).*?(?:in|of).*?(?:Computer Science|Engineering|Business|Science)',
            r'(?:Education|EDUCATION)(?:[\n\r].+){1,10}'
        ]
        
        for pattern in education_patterns:
            education_match = re.search(pattern, resume_text, re.DOTALL | re.IGNORECASE)
            if education_match:
                edu_text = education_match.group(1).strip() if len(education_match.groups()) > 0 else education_match.group(0).strip()
                if len(edu_text) > 5:
                    info['education'] = edu_text
                    break
                    
        if 'education' not in info or not info['education']:
            info['education'] = "Bachelor's Degree in Computer Science"
            
        experience_patterns = [
            regex_config.work_experience_regex,
            r'(?:EXPERIENCE|WORK EXPERIENCE|EMPLOYMENT).*?\n(.*?)(?=\n\s*\n|\n[A-Z\s]{2,}|\Z)',
            r'(?:19|20)\d{2}\s*[-–—]\s*(?:19|20)\d{2}|(?:19|20)\d{2}\s*[-–—]\s*Present',
            r'(?:Company|Employer|Job Title|Position):\s*.*?\n',
            r'(?:Experience|EXPERIENCE|Work Experience|WORK EXPERIENCE)(?:[\n\r].+){1,20}'
        ]
        
        for pattern in experience_patterns:
            experience_match = re.search(pattern, resume_text, re.DOTALL | re.IGNORECASE)
            if experience_match:
                exp_text = experience_match.group(1).strip() if len(experience_match.groups()) > 0 else experience_match.group(0).strip()
                if len(exp_text) > 10:
                    info['work_experience'] = exp_text
                    break
                    
        if 'work_experience' not in info or not info['work_experience']:
            info['work_experience'] = "Software Engineer with 2+ years of experience"
            
        project_patterns = [
            regex_config.projects_regex,
            r'(?:PROJECTS|PROJECT EXPERIENCE|PERSONAL PROJECTS|UNIVERSITY PROJECTS).*?\n(.*?)(?=\n\s*\n|\n[A-Z\s]{2,}|\Z)',
            r'(?:Developed|Created|Built|Implemented|Designed).*?(?:application|system|website|solution)',
            r'(?:Projects|PROJECTS|Project Experience|PROJECT EXPERIENCE)(?:[\n\r].+){1,15}'
        ]
        
        for pattern in project_patterns:
            project_match = re.search(pattern, resume_text, re.DOTALL | re.IGNORECASE)
            if project_match:
                proj_text = project_match.group(1).strip() if len(project_match.groups()) > 0 else project_match.group(0).strip()
                if len(proj_text) > 10:
                    info['projects'] = proj_text
                    break
                    
        if 'projects' not in info or not info['projects']:
            info['projects'] = "Various projects developing software applications"
            
        activity_patterns = [
            regex_config.activities_regex,
            r'(?:ACTIVITIES|LEADERSHIP|VOLUNTEER|EXTRACURRICULAR|ADDITIONAL).*?\n(.*?)(?=\n\s*\n|\n[A-Z\s]{2,}|\Z)',
            r'(?:participated|organized|volunteered|member of|founded)',
            r'(?:Activities|ACTIVITIES|Leadership|LEADERSHIP|Volunteer|VOLUNTEER)(?:[\n\r].+){1,10}'
        ]
        
        for pattern in activity_patterns:
            activity_match = re.search(pattern, resume_text, re.DOTALL | re.IGNORECASE)
            if activity_match:
                act_text = activity_match.group(1).strip() if len(activity_match.groups()) > 0 else activity_match.group(0).strip()
                if len(act_text) > 5:
                    info['activities'] = act_text
                    break
                    
        if 'activities' not in info or not info['activities']:
            info['activities'] = "Participation in professional development activities"
            
        skills_patterns = [
            r'(?:SKILLS|TECHNICAL SKILLS|TECHNOLOGIES|TOOLS).*?\n(.*?)(?=\n\s*\n|\n[A-Z\s]{2,}|\Z)',
            r'(?:Technical Skills|Technical):\s*(.*?)(?=\n\n|\n[A-Z\s]|\Z)',
            r'(?:Proficient in|Skilled in|Knowledge of|Experience with).*?(?:languages|frameworks|technologies|tools)',
            r'(?:Programming Languages|Software|Tools):\s*(.*?)(?=\n\n|\Z)', 
            r'(?:Skills|SKILLS|Technical Skills|TECHNICAL SKILLS)(?:[\n\r].+){1,10}',
            r'(?:ADDITIONAL).*?Technical Skills:\s*(.*?)(?=\n\n|\n[A-Z\s]|\Z)'
        ]
        
        tech_skills = []
        
        for pattern in skills_patterns:
            skills_match = re.search(pattern, resume_text, re.DOTALL | re.IGNORECASE)
            if skills_match:
                if len(skills_match.groups()) > 0:
                    skills_text = skills_match.group(1).strip()
                else:
                    skills_text = skills_match.group(0).strip()
                
                if len(skills_text) > 5:
                    info['skills'] = skills_text
                    break
        
        if 'skills' not in info or len(info['skills']) < 10:
            common_tech = [
                "Python", "Java", "C\\+\\+", "JavaScript", "HTML", "CSS", "SQL", "R", 
                "Tableau", "Excel", "PowerBI", "TensorFlow", "PyTorch", "Keras", 
                "Machine Learning", "Deep Learning", "NLP", "Data Analysis",
                "React", "Angular", "Vue", "Node.js", "Django", "Flask", "Spring",
                "AWS", "Azure", "GCP", "Docker", "Kubernetes", "Linux", "PHP",
                "MATLAB"
            ]
            
            for tech in common_tech:
                if re.search(r'\b' + tech + r'\b', resume_text, re.IGNORECASE):
                    tech_skills.append(tech)
            
            if tech_skills:
                if 'skills' in info and len(info['skills']) > 5:
                    info['skills'] += ", " + ", ".join(tech_skills)
                else:
                    info['skills'] = ", ".join(tech_skills)
        
        if 'skills' not in info or not info['skills'] or len(info['skills']) < 10:
            info['skills'] = role_default_skills.get(role, f"Skills relevant to {role}")
            
        summary = f"Professional with experience in {role} related fields."
        
        if role in info.get('skills', ''):
            summary += f" Skilled in {info['skills']}."
        elif info.get('skills') and info.get('skills') != role_default_skills.get(role, f"Skills relevant to {role}"):
            summary += f" Skilled in {info['skills']}."
            
        experience_keywords = []
        for pattern in [role, "developed", "created", "managed", "implemented", "designed"]:
            if pattern.lower() in info.get('work_experience', '').lower():
                experience_keywords.append(pattern)
                
        if experience_keywords:
            summary += f" Experience includes {', '.join(experience_keywords[:3])}."
            
        education_level = ""
        for degree in ["PhD", "Master", "Bachelor"]:
            if degree.lower() in info.get('education', '').lower():
                education_level = degree
                break
                
        if education_level:
            summary += f" Holds a {education_level}'s degree."
            
        if len(summary) < 100:
            role_summaries = {
                "Data Scientist": " Passionate about extracting insights from complex datasets to drive business decisions. Experienced in statistical modeling and machine learning algorithms.",
                "AI Engineer": " Focuses on building intelligent systems that solve real-world problems. Stays current with the latest advancements in artificial intelligence.",
                "Software Engineer": " Committed to writing clean, maintainable code following best practices. Experienced in full software development lifecycle.",
                "Web Developer": " Creates responsive and intuitive web applications with focus on user experience. Skilled in frontend and backend technologies.",
                "DevOps Engineer": " Ensures smooth deployment pipelines and robust infrastructure. Experience with cloud platforms and containerization technologies."
            }
            summary += role_summaries.get(role, " Demonstrates strong technical capabilities and problem-solving skills.")
            
        info['summary'] = summary
        
        print("\nInformación extraída del CV:")
        for key, value in info.items():
            print(f"{key}: {value[:100]}...")

        return info